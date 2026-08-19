from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Trading Bot Setup Instructions', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Follow these exact steps to deploy your live trading bot on a new computer.')

# Phase 1
doc.add_heading('Phase 1: Install the Requirements', level=1)
doc.add_paragraph('1. Download Python: Install Python 3.10+ from python.org (Make sure to check the "Add Python to PATH" box during installation).')
doc.add_paragraph('2. Download Git: Install Git from git-scm.com.')
doc.add_paragraph('3. Install Interactive Brokers: Download and install IBKR Trader Workstation (TWS) or IB Gateway on the new computer.')

# Phase 2
doc.add_heading('Phase 2: Setup Interactive Brokers', level=1)
doc.add_paragraph('1. Log into TWS or Gateway on the new computer.')
doc.add_paragraph('2. Go to Settings > API > Settings.')
doc.add_paragraph('3. Check "Enable ActiveX and Socket Clients".')
doc.add_paragraph('4. Set the Socket Port to 7497 (for Paper Trading) or 7496 (for Live Trading).')
doc.add_paragraph('5. Uncheck "Read-Only API".')

# Phase 3
doc.add_heading('Phase 3: Download and Run the Bot', level=1)
doc.add_paragraph('Open your terminal (PowerShell) on the new computer and run these exact commands in order:')

p = doc.add_paragraph()
run = p.add_run('1. Download the code from GitHub:\n')
run.bold = True
code1 = p.add_run('git clone https://github.com/abhims2012/IBKR.git\ncd IBKR')
code1.font.name = 'Courier New'

p2 = doc.add_paragraph()
run2 = p2.add_run('2. Setup the Python Environment:\n')
run2.bold = True
code2 = p2.add_run('python -m venv venv\n.\\venv\\Scripts\\activate\npip install -r requirements.txt')
code2.font.name = 'Courier New'

p3 = doc.add_paragraph()
run3 = p3.add_run('3. Start the Bot:\n')
run3.bold = True
code3 = p3.add_run('.\\run_forever.ps1')
code3.font.name = 'Courier New'

doc.add_paragraph('\nThat\'s it! The bot will instantly connect to your IBKR software on the new computer, read your watchlist, and pick up trading right where it left off, while automatically updating your live dashboard.')

doc.save('Bot_Setup_Instructions.docx')
print("Word Document created successfully!")
